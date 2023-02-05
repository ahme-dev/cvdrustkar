import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


#class struct
class experience:
	def __init__(self, years,place,about):
		self.years = years
		self.place = place
		self.about = about
class education:
	def __init__(self, year,degree,place,about):
		self.year = year
		self.place = place
		self.degree = degree
		self.about = about
class skill:
	def __init__(self, name,about):
		self.name = name
		self.about = about


def make_file(directory, filename, colours, info_list, exp_list, edu_list, ski_list):
	# make file and set colours
	cv = docx.Document()
	col1, col2, col3 = set_colours(colours)

	# add basic info
	add_name(cv, info_list[0], col1)
	add_info(cv, info_list[1], info_list[2], info_list[3], col3)

	# add experience
	if exp_list:
		# add head for section
		h1 = "ئەزموون و کارەکان"
		add_head(cv, h1, col1)
		# add each experience as body
		for i in exp_list:
			add_body_experience(cv, i.years, i.place, i.about, col2, col3)

	# add education
	if edu_list:
		# add head for section
		h2 = "بڕوانامەکان"
		add_head(cv, h2, col1)
		# add each education as body
		for i in edu_list:
			add_body_education(cv, i.year, i.degree, i.place, i.about, col2, col3)

	# add skills
	if ski_list:
		# add head for section
		h3 = "بەهرە و تواناکان"
		add_head(cv, h3, col1)
		# add each skill as body
		for i in ski_list:
			add_body_skill(cv, i.name, i.about, col2, col3)

	# save the file in rand url
	cv.save(directory + filename)


# return a list of three color values
def set_colours(colour):
	if colour == "blue":
		colour1 = RGBColor(0x50, 0x70, 0x90)
		colour2 = RGBColor(0x20, 0x30, 0x40)
	elif colour == "red":
		colour1 = RGBColor(0x90, 0x50, 0x50)
		colour2 = RGBColor(0x40, 0x20, 0x20)
	elif colour == "brown":
		colour1 = RGBColor(0x90, 0x70, 0x50)
		colour2 = RGBColor(0x40, 0x30, 0x20)
	elif colour == "green":
		colour1 = RGBColor(0x50, 0x90, 0x50)
		colour2 = RGBColor(0x20, 0x40, 0x20)
	colour3 = RGBColor(0x20, 0x20, 0x20)
	return colour1, colour2, colour3


def add_name(doc, name, colour):
	if name:
		title = doc.add_heading("", 0)
		title.add_run(name).font.color.rgb = colour
		title.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# add other information
def add_info(doc, address, number, email, colour):
	if address or number or email:
		info = doc.add_heading("", 5)
		info.alignment = WD_ALIGN_PARAGRAPH.LEFT
		info.paragraph_format.left_indent = Pt(35)
		info.paragraph_format.right_indent = Pt(70)
		if address:
			info.add_run(address).font.color.rgb = colour
			info.add_run("\n")
		if number:
			info.add_run(number).font.color.rgb = colour
			info.add_run("\n")
		if email:
			info.add_run(email).font.color.rgb = colour
			info.add_run("\n")


# add header for a section
def add_head(doc, title, colour):
	head = doc.add_heading("", 1)
	head.add_run(title).font.color.rgb = colour
	head.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_body_experience(doc, years, workplace, about, colour1, colour2):
	if years or workplace:
		exp_a = doc.add_heading("", 5)
		exp_a.add_run(workplace).font.color.rgb = colour1
		exp_a.add_run("   ")
		exp_a.add_run(years).font.color.rgb = colour1
		exp_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	if about:
		exp_b = doc.add_paragraph()
		exp_b.add_run(about).font.color.rgb = colour2
		exp_b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		exp_b.paragraph_format.right_indent = Pt(70)
		exp_b.paragraph_format.left_indent = Pt(70)


def add_body_education(doc, years, degree, school, about, colour1, colour2):
	if years or school or degree:
		edu_a = doc.add_heading("", 5)
		edu_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		if school:
			edu_a.add_run(school).font.color.rgb = colour1
			edu_a.add_run(" - ")
		if degree:
			edu_a.add_run(degree).font.color.rgb = colour1
			edu_a.add_run("   ")
		if years:
			edu_a.add_run(years).font.color.rgb = colour1

	if about:
		edu_b = doc.add_paragraph()
		edu_b.add_run(about).font.color.rgb = colour2
		edu_b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		edu_b.paragraph_format.right_indent = Pt(35)
		edu_b.paragraph_format.left_indent = Pt(70)


def add_body_skill(doc, skill, about, colour1, colour2):
	if skill:
		ski = doc.add_paragraph()
		ski.add_run(skill).font.color.rgb = colour1

		ski.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		ski.paragraph_format.right_indent = Pt(35)
		ski.paragraph_format.left_indent = Pt(70)
		ski.paragraph_format.space_before = Pt(10)

	if about:
		ski.add_run(": ").font.color.rgb = colour1
		ski.add_run(about).font.color.rgb = colour2
